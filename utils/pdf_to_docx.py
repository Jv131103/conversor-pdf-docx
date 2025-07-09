import io
import os

import fitz
from docx import Document
from docx.shared import Inches
from pdf2docx import Converter
from PIL import Image


class ConvertPdfToDocx:
    def __init__(
            self,
            pdf_file: str = "input.pdf",
            name_file_docx: str = "output.docx"
    ) -> None:
        self.__file_pdf = pdf_file
        self.__name_file_docx = name_file_docx
        self._conversor = Converter(self.__file_pdf)

    def __convert_to_docx(self) -> bool:
        try:
            if not self.is_valid_pdf():
                print(
                    "[-] 'PDF' inválido.",
                    "Não foi possível realizar conversão para 'DOCX'")
                return False

            self._conversor.convert(self.__name_file_docx)
            return True
        except Exception as e:
            print(f"[-] Erro interno ao converter 'PDF' para 'DOCX': {e}")
            return False
        finally:
            self._conversor.close()

    def is_valid_pdf(self) -> bool:
        return (
            self.__file_pdf.lower().endswith('.pdf')
            and os.path.isfile(self.__file_pdf)
        )

    def __extract_image_for_docx(self) -> bool:
        try:
            if not self.is_valid_pdf():
                print("[-] PDF inválido. Não foi possível extrair imagens.")
                return False

            doc = fitz.open(self.__file_pdf)
            word_doc = (
                Document(self.__name_file_docx)
                if os.path.exists(self.__name_file_docx)
                else Document()
            )

            for i, page in enumerate(doc):
                imagens = page.get_images(full=True)

                for img_index, img in enumerate(imagens):
                    try:
                        xref = img[0]
                        base_img = doc.extract_image(xref)
                        image_bytes = base_img["image"]
                        image_stream = io.BytesIO(image_bytes)

                        # Testa se é uma imagem válida
                        Image.open(image_stream)
                        image_stream.seek(0)

                        word_doc.add_picture(image_stream, width=Inches(5))
                        word_doc.add_paragraph(
                            f"Página {i + 1} - Imagem {img_index + 1}"
                        )
                    except Exception as e:
                        print(
                            f"[-] Imagem com problema na página {i + 1}: {e}"
                        )

            word_doc.save(self.__name_file_docx)
            print("[+] Imagens inseridas no DOCX com sucesso!")
            return True

        except Exception as e:
            print(f"[-] Erro ao extrair imagens do PDF: {e}")
            return False

    def __is_real_text(self) -> bool:
        try:
            with fitz.open(self.__file_pdf) as doc:
                for page in doc:
                    if page.get_text().strip():
                        return True
            return False
        except Exception as e:
            print(f"[-] Erro ao verificar texto no PDF: {e}")
            return False

    def execute_conversor(self) -> bool:
        if not self.is_valid_pdf():
            print("[-] Arquivo PDF inválido.")
            return False

        print("[*] Verificando se o PDF contém texto real...")
        if self.__is_real_text():
            print("[+] Texto detectado. Convertendo com pdf2docx...")
            return self.__convert_to_docx()
        else:
            print("[*] Nenhum texto detectado. Extraindo imagens...")
            return self.__extract_image_for_docx()


if __name__ == "__main__":
    conversor = ConvertPdfToDocx(
        "estruturas-de-dados-udemy.pdf", "meu_docx.docx"
    )
    conversor.execute_conversor()
